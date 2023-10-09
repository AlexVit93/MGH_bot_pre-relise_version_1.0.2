import os
from docx import Document
from uuid import uuid4
from google.oauth2.service_account import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload


def create_docx(data):
    doc = Document()
    doc.add_paragraph(f"User ID: {data['user_id']}")
    doc.add_paragraph(f"Name: {data['name']}")
    doc.add_paragraph(f"Phone_number: {data['phone_number']}")
    doc.add_paragraph(f"Gender: {data['gender']}")
    doc.add_paragraph(f"Age: {data['age']}")

    for question, answer in {**data["child_answers"], **data["answers"]}.items():
        doc.add_paragraph(f"{question}: {answer}")

    doc.add_paragraph(f"Recommendations: {data['recommendations']}")

    filename = f"{uuid4().hex}.docx"
    doc.save(filename)

    return filename


def upload_to_drive(filename):
    creds = Credentials.from_service_account_file("local-volt-399322-5c1d9c000102.json")
    drive_service = build("drive", "v3", credentials=creds)

    file_metadata = {"name": filename, "parents": ["1Ly8tlk0aCEEql5_cN7_EnO8H_rQ9Shn0"]}
    media = MediaFileUpload(filename)
    drive_service.files().create(body=file_metadata, media_body=media).execute()


def generate_and_upload(data):
    filename = create_docx(data)
    upload_to_drive(filename)
    os.remove(filename)
